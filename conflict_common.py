# import streamlit as st
# import fitz
# from io import BytesIO
# import pandas as pd
# from fuzzywuzzy import fuzz  # For fuzzy string matching
# import time  # To track execution time
# import re  # For extracting years
# import datetime

# # Function to normalize publication titles
# def clean_title(title):
#     return " ".join(title.lower().strip().replace("-", " ").split())  # Normalize hyphens, spaces, and case

# # Function to strip .pdf extension from file name
# def clean_filename(file_name):
#     return file_name.rsplit('.', 1)[0]  # Remove the .pdf extension

# # Function to check if the text color is blue
# def is_blue(color_int):
#     r = (color_int >> 16) & 0xFF
#     g = (color_int >> 8) & 0xFF
#     b = color_int & 0xFF
#     return b > r and b > g

# # Function to extract titles and years from PDF
# def extract_titles_and_years(file_object):
#     titles_and_years = []

#     if file_object is None or file_object.getbuffer().nbytes == 0:
#         return []

#     file_object.seek(0)
#     pdf_document = fitz.open(stream=file_object.read(), filetype="pdf")

#     try:
#         for page_number in range(len(pdf_document)):
#             page = pdf_document.load_page(page_number)
#             current_title = ""
#             current_year = ""
#             for text_instance in page.get_text("dict")["blocks"]:
#                 if "lines" in text_instance:
#                     for line in text_instance["lines"]:
#                         for span in line["spans"]:
#                             # Extract titles (blue text)
#                             if is_blue(span['color']):
#                                 current_title += span['text'] + " "
#                             else:
#                                 # Check if the text is a year (check for 4-digit number)
#                                 year_match = re.match(r"\b(19|20)\d{2}\b", span['text'])
#                                 if year_match:
#                                     current_year = year_match.group(0)
#                                     if current_title.strip() and current_year:
#                                         titles_and_years.append((clean_title(current_title.strip()), current_year))
#                                         current_title = ""  # Reset current title after saving
#                                         current_year = ""  # Reset year after saving
#     finally:
#         pdf_document.close()

#     return titles_and_years

# # Function to find similar titles between two sets using fuzzy matching
# def find_similar_titles(set1, set2, threshold=95):
#     matches = set()
#     for title1 in set1:
#         for title2 in set2:
#             if fuzz.ratio(title1, title2) >= threshold:
#                 matches.add(title1)  # Store only one version
#     return matches

# # Streamlit UI
# st.title("ConflictMatch: Detecting conflicts through common publications")

# # Add the attribution text below the title
# st.markdown(
#     '"This ConflictMatch tool is implemented based on the work done by Rohit Ramachandran, Rutgers University, NJ, USA"'
# )

# # Add instructions banner
# st.markdown("""
# ### Instructions:
# 1. Upload at least 1 Google Scholar PDF in each group. Ensure all the papers are printed to PDF. See 'show more' bottom of google scholar link.
# 2. The algorithm will check for common publications across these group of researchers for different years selected and identify conflicts.
# 3. You can also view the tables showing common publications and corresponding conflicts.  
# 4. Once the analysis is complete, you can download the results as a CSV file.
# """)


# # Get the current year dynamically
# current_year = datetime.datetime.now().year

# # Dropdown for year filtering
# year_options = ["All Years"] + list(range(current_year, 1950, -1))  # Start from current year
# selected_years = st.multiselect("Select years to include:", year_options, default=year_options[1:6])  # Default: last 5 years

# # If "All Years" is selected, include all years
# if "All Years" in selected_years:
#     selected_years = set(map(str, range(current_year, 1950, -1)))  # Convert all years to string
# else:
#     selected_years = set(map(str, [year for year in selected_years if year != "All Years"]))  # Convert selected years to string


# # File uploader for Group 1 and Group 2
# group1_files = st.file_uploader("Upload Google Scholar PDFs for Group 1", type=["pdf"], accept_multiple_files=True)
# group2_files = st.file_uploader("Upload Google Scholar PDFs for Group 2", type=["pdf"], accept_multiple_files=True)

# group1_uploaded_names = [clean_filename(file.name) for file in group1_files] if group1_files else []
# group2_uploaded_names = [clean_filename(file.name) for file in group2_files] if group2_files else []
# # Run simulation button
# run_simulation_button = st.button("Run Simulation")

# if run_simulation_button:
#     # Record start time
#     start_time = time.time()

#     # Process Group 1 files
#     group1_data = {}
#     for file in group1_files:
#         researcher_name = clean_filename(file.name)
#         extracted_data = extract_titles_and_years(file)
#         filtered_titles = {title for title, year in extracted_data if year in selected_years}
#         group1_data[researcher_name] = filtered_titles

#     # Process Group 2 files
#     group2_data = {}
#     for file in group2_files:
#         researcher_name = clean_filename(file.name)
#         extracted_data = extract_titles_and_years(file)
#         filtered_titles = {title for title, year in extracted_data if year in selected_years}
#         group2_data[researcher_name] = filtered_titles

#     # Ensure we have data from both groups
#     if group1_data and group2_data:
#         st.subheader("Common Publications Between Researchers from Group 1 and Group 2")
#         group1_names = [name for name in group1_uploaded_names if name in group1_data]
#         group2_names = [name for name in group2_uploaded_names if name in group2_data]
#         # Create a matrix for comparisons
#         comparison_matrix = []
#         group1_names = list(group1_data.keys())
#         group2_names = list(group2_data.keys())

#         for researcher1 in group1_names:
#             row = []
#             titles1 = group1_data[researcher1]
#             for researcher2 in group2_names:
#                 titles2 = group2_data[researcher2]
#                 common_titles = find_similar_titles(titles1, titles2, threshold=95)
#                 row.append(len(common_titles))  # Store the number of common publications
#             comparison_matrix.append(row)

#         # Create the DataFrame for the matrix with filenames as labels
#         df_comparisons = pd.DataFrame(comparison_matrix, index=[f"{i+1}. Grp 1: {name}" for i, name in enumerate(group1_names)], 
#                                       columns=[f"Grp 2: {name}" for name in group2_names])

#         # Display the comparison matrix with headers
#         st.dataframe(df_comparisons.style.set_properties(**{'text-align': 'center', 'white-space': 'pre-wrap'}))

#         # Convert to COI matrix
#         df_coi = df_comparisons.applymap(lambda x: "COI" if x > 0 else "")

#         # Style COI matrix
#         def highlight_coi(val):
#             return "background-color: #FFCCCC; font-weight: bold;" if val == "COI" else ""

#         st.subheader("Conflict of Interest (COI) Matrix")
#         st.dataframe(df_coi.style.applymap(highlight_coi))

#         # Allow download of the CSV file
#         @st.cache_data
#         def convert_df(df):
#             return df.to_csv(index=True).encode('utf-8')

#         csv_data = convert_df(df_comparisons)
#         st.download_button(
#             label="Download Common publications as CSV",
#             data=csv_data,
#             file_name="publication_comparisons_matrix.csv",
#             mime="text/csv"
#         )

#         # Table 2: Download COI Matrix CSV (df_coi)
#         csv_data_coi = convert_df(df_coi)
#         st.download_button(
#         label="Download COI Matrix as CSV",
#         data=csv_data_coi,
#         file_name="publication_coi_matrix.csv",
#         mime="text/csv"
#         )

#     # Record end time and compute total simulation time
#     end_time = time.time()
#     total_time = (end_time - start_time) / 60  # Convert seconds to minutes

#     # Display the total simulation time
#     st.subheader(f"Total simulation time: {total_time:.2f} minutes")

import streamlit as st
import pymupdf
from io import BytesIO
import pandas as pd
from fuzzywuzzy import fuzz
import time
import re
import datetime
import io 
from docx import Document

# === Utility Functions ===

def clean_title(title):
    return " ".join(title.lower().strip().replace("-", " ").split())

def clean_filename(file_name):
    return file_name.rsplit('.', 1)[0]

def is_blue(color_int):
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return b > r and b > g

def extract_titles_and_years(file_object):
    titles_and_years = []
    if file_object is None or file_object.getbuffer().nbytes == 0:
        return []

    file_object.seek(0)
    pdf_document = pymupdf.open(stream=file_object.read(), filetype="pdf")

    try:
        for page in pdf_document:
            current_title = ""
            current_year = ""
            for block in page.get_text("dict")["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            if is_blue(span['color']):
                                current_title += span['text'] + " "
                            else:
                                year_match = re.match(r"\b(19|20)\d{2}\b", span['text'])
                                if year_match:
                                    current_year = year_match.group(0)
                                    if current_title.strip() and current_year:
                                        titles_and_years.append((clean_title(current_title.strip()), current_year))
                                        current_title = ""
                                        current_year = ""
    finally:
        pdf_document.close()

    return titles_and_years


#-----------------
# def extract_institution_names(file_object):
#     institutions = set()
#     if file_object is None or file_object.getbuffer().nbytes == 0:
#         return institutions

#     file_object.seek(0)
#     pdf_document = pymupdf.open(stream=file_object.read(), filetype="pdf")

#     try:
#         for page in pdf_document:
#             text = page.get_text("text")
#             lines = text.splitlines()

#             # Check last 10 lines for a "page 1" marker
#             last_lines = lines[-10:]
#             is_page_one = any(
#                 re.search(r'\b1\s*/\s*\d+\b', line.lower()) or
#                 re.search(r'\b1\s+of\s+\d+\b', line.lower())
#                 for line in last_lines
#             )

#             if is_page_one:
#                 for line in lines[:20]:
#                     normalized_line = line.strip().lower()
#                     normalized_line = re.sub(r'(?<=\S)-(?=\S)', ' ', normalized_line)
#                     normalized_line = re.sub(r'[\s\-–—]+', ' ', normalized_line)
#                     normalized_line = normalized_line.replace('-', '').replace('–', '').replace('—', '')
#                     normalized_line = normalized_line.replace("univeristy", "university")
#                     normalized_line = normalized_line.replace("colledge", "college")
#                     normalized_line = normalized_line.replace("centeres", "centers")
#                     normalized_line = normalized_line.replace("scool", "school")
#                     normalized_line = normalized_line.replace("instutute", "institute")
#                     normalized_line = normalized_line.replace("laborotory", "laboratory")
                    
#                     # Remove any trailing commas or periods from the word
#                     normalized_line = re.sub(r'[,.]+$', '', normalized_line)

#                     if re.search(r'\b(uni|univ|university|college|institute|school|centre|center|laboratory|co|company)\b', normalized_line):
#                         institutions.add(normalized_line)
#     finally:
#         pdf_document.close()

#     return institutions


#---------- original
# def extract_institution_names(file_object):
#     institutions = set()
#     if file_object is None or file_object.getbuffer().nbytes == 0:
#         return institutions

#     alias_mapping = {
#         "rutgers": "rutgers university",
#         "mit": "massachusetts institute of technology",
#         "caltech": "california institute of technology",
#         "ucb": "university of california berkeley",
#         "ucla": "university of california los angeles",
#         "ucsd": "university of california san diego",
#         "ucsf": "university of california san francisco",
#         "cmu": "carnegie mellon university",
#         "ntu": "nanyang technological university",
#         "nus": "national university of singapore",
#         "tsinghua": "tsinghua university",
#         "epfl": "École polytechnique fédérale de lausanne",
#         "kaist": "korea advanced institute of science and technology",
#         "kth": "kth royal institute of technology",
#         "rwth": "rwth aachen university",
#         "oxford": "university of oxford",
#         "cambridge": "university of cambridge",
#         "harvard": "harvard university",
#         "stanford": "stanford university",
#         "princeton": "princeton university",
#         "yale": "yale university",
#         "columbia": "columbia university",
#         "upenn": "university of pennsylvania",
#         "uchicago": "university of chicago",
#         "tokyo tech": "tokyo institute of technology",
#         "uofm": "university of michigan",
#         "utexas": "university of texas",
#         "nyu": "new york university",
#         "ncsu": "north carolina state university",
#         "wustl": "washington university in st. louis",
#         "duke": "duke university",
#         "cornell": "cornell university",
#         "berkeley": "university of california berkeley",
#         "brown": "brown university",
#         "purdue": "purdue university",
#         "georgia tech": "georgia institute of technology",
#         "uci": "university of california irvine"
#     }

#     stop_words = {"and", "or", "with", "at", "in", "of"}

#     file_object.seek(0)
#     pdf_document = pymupdf.open(stream=file_object.read(), filetype="pdf")

#     try:
#         for page in pdf_document:
#             text = page.get_text("text")
#             lines = text.splitlines()

#             last_lines = lines[-10:]
#             is_page_one = any(
#                 re.search(r'\b1\s*/\s*\d+\b', line.lower()) or
#                 re.search(r'\b1\s+of\s+\d+\b', line.lower())
#                 for line in last_lines
#             )

#             if is_page_one:
#                 for line in lines[:20]:
#                     normalized_line = line.strip().lower()
#                     normalized_line = re.sub(r'(?<=\S)-(?=\S)', ' ', normalized_line)
#                     normalized_line = re.sub(r'[\s\-–—]+', ' ', normalized_line)
#                     normalized_line = normalized_line.replace("univeristy", "university")
#                     normalized_line = normalized_line.replace("colledge", "college")
#                     normalized_line = normalized_line.replace("centeres", "centers")
#                     normalized_line = normalized_line.replace("scool", "school")
#                     normalized_line = normalized_line.replace("instutute", "institute")
#                     normalized_line = normalized_line.replace("laborotory", "laboratory")
#                     normalized_line = re.sub(r'[,.]+$', '', normalized_line)
#                     normalized_line = normalized_line.replace(",", "")

#                     # Remove leading stop words
#                     tokens = normalized_line.split()
#                     while tokens and tokens[0] in stop_words:
#                         tokens.pop(0)
#                     cleaned_line = " ".join(tokens)

#                     if re.search(r'\b(uni|univ|university|college|institute|school|centre|center|laboratory|co|company)\b', cleaned_line) or cleaned_line in alias_mapping:
#                         full_name = alias_mapping.get(cleaned_line, cleaned_line)
#                         institutions.add(full_name)
#     finally:
#         pdf_document.close()

#     return institutions

 
# def extract_institution_names(file_object):
#     institutions = set()
#     if not file_object or file_object.getbuffer().nbytes == 0:
#         return institutions

#     file_object.seek(0)
#     pdf = pymupdf.open(stream=file_object.read(), filetype="pdf")
#     try:
#         # grab plain text from page 1
#         text = pdf[0].get_text("text")
#         lines = text.splitlines()[:50]  # look at first 50 lines

#         # 1) First pass: title‐cased regex (captures e.g. "Rutgers University")
#         titlecased = re.compile(r'([A-Z][\w&\.,\s]{2,30}\b(?:University|College|Institute)\b)')
#         for line in lines:
#             m = titlecased.search(line)
#             if m:
#                 institutions.add(m.group(1))
#         if institutions:
#             return institutions

#         # 2) Fallback: your original lowercase, typo‐fixing logic
#         for line in lines:
#             nl = line.strip().lower()
#             # normalize hyphens/dashes and fix typos
#             nl = re.sub(r'(?<=\S)-(?=\S)', ' ', nl)
#             nl = re.sub(r'[\s\-–—]+',' ', nl)
#             for typo, fix in [
#                 ("univeristy","university"),("colledge","college"),
#                 ("centeres","centers"),("scool","school"),
#                 ("instutute","institute"),("laborotory","laboratory")
#             ]:
#                 nl = nl.replace(typo, fix)
#             # if it still contains a keyword, add it
#             if re.search(r'\b(univ|university|college|institute|school|centre|center|laboratory)\b', nl):
#                 institutions.add(nl.title())

#     finally:
#         pdf.close()

#     return institutions

### Works

# def extract_institution_names(file_object): 
#     institutions = set()
#     if file_object is None or file_object.getbuffer().nbytes == 0:
#         return institutions

#     alias_mapping = {
#         "rutgers": "rutgers university",
#         "mit": "massachusetts institute of technology",
#         "caltech": "california institute of technology",
#         "ucb": "university of california berkeley",
#         "ucla": "university of california los angeles",
#         "ucsd": "university of california san diego",
#         "ucsf": "university of california san francisco",
#         "cmu": "carnegie mellon university",
#         "ntu": "nanyang technological university",
#         "nus": "national university of singapore",
#         "tsinghua": "tsinghua university",
#         "epfl": "École polytechnique fédérale de lausanne",
#         "kaist": "korea advanced institute of science and technology",
#         "kth": "kth royal institute of technology",
#         "rwth": "rwth aachen university",
#         "oxford": "university of oxford",
#         "cambridge": "university of cambridge",
#         "harvard": "harvard university",
#         "stanford": "stanford university",
#         "princeton": "princeton university",
#         "yale": "yale university",
#         "columbia": "columbia university",
#         "upenn": "university of pennsylvania",
#         "uchicago": "university of chicago",
#         "tokyo tech": "tokyo institute of technology",
#         "uofm": "university of michigan",
#         "utexas": "university of texas",
#         "nyu": "new york university",
#         "ncsu": "north carolina state university",
#         "wustl": "washington university in st. louis",
#         "duke": "duke university",
#         "cornell": "cornell university",
#         "berkeley": "university of california berkeley",
#         "brown": "brown university",
#         "purdue": "purdue university",
#         "georgia tech": "georgia institute of technology",
#         "uci": "university of california irvine"
#     }
#     stop_words = {"and", "or", "with", "at", "in", "of"}

#     # regex to catch Title‑Cased names ending in University/College/etc.
#     title_regex = re.compile(
#         r'([A-Z][\w&\.\,\s]{2,40}\b(?:University|College|Institute|School|Centre|Center|Laboratory)\b)'
#     )

#     file_object.seek(0)
#     pdf = pymupdf.open(stream=file_object.read(), filetype="pdf")
#     try:
#         # only look at page 1
#         page = pdf[0]
#         lines = page.get_text("text").splitlines()
#         header = lines[:30]  # extend window a bit

#         # 1) Title‐cased pass
#         for L in header:
#             m = title_regex.search(L)
#             if m:
#                 institutions.add(m.group(1).strip())
#         if institutions:
#             return institutions

#         # 2) Fallback: lowercase + typo fixes + alias + keyword
#         for L in header:
#             nl = L.strip().lower()
#             # normalize hyphens/dashes
#             nl = re.sub(r'(?<=\S)-(?=\S)', ' ', nl)
#             nl = re.sub(r'[\s\-–—]+', ' ', nl)
#             # fix common typos
#             for typo, fix in [
#                 ("univeristy","university"),
#                 ("colledge","college"),
#                 ("centeres","centers"),
#                 ("scool","school"),
#                 ("instutute","institute"),
#                 ("laborotory","laboratory")
#             ]:
#                 nl = nl.replace(typo, fix)
#             # strip trailing punctuation
#             nl = re.sub(r'[,\.;:]+$', '', nl)

#             # drop leading stop‑words
#             tokens = nl.split()
#             while tokens and tokens[0] in stop_words:
#                 tokens.pop(0)
#             cleaned = " ".join(tokens)

#             # 2a) alias substring match
#             for key, full in alias_mapping.items():
#                 if key in cleaned.split():  # match on a whole word
#                     institutions.add(full)
#             # 2b) keyword match
#             if re.search(
#                 r'\b(univ|university|college|institute|school|centre|center|laboratory|co|company)\b',
#                 cleaned
#             ):
#                 # title‑case for presentation
#                 institutions.add(cleaned.title())

#         return institutions

#     finally:
#         pdf.close()

import re
import pymupdf

def extract_institution_names(file_object):
    institutions = set()
    if not file_object or file_object.getbuffer().nbytes == 0:
        return institutions

    # 1) your "short key" → full-name map
    alias_mapping = {
        "rutgers": "rutgers university",
        "mit": "massachusetts institute of technology",
        "caltech": "california institute of technology",
        "ucb": "university of california berkeley",
        "ucla": "university of california los angeles",
        "ucsd": "university of california san diego",
        "ucsf": "university of california san francisco",
        "cmu": "carnegie mellon university",
        "ntu": "nanyang technological university",
        "nus": "national university of singapore",
        "tsinghua": "tsinghua university",
        "epfl": "École polytechnique fédérale de lausanne",
        "kaist": "korea advanced institute of science and technology",
        "kth": "kth royal institute of technology",
        "rwth": "rwth aachen university",
        "oxford": "university of oxford",
        "cambridge": "university of cambridge",
        "harvard": "harvard university",
        "stanford": "stanford university",
        "princeton": "princeton university",
        "yale": "yale university",
        "columbia": "columbia university",
        "upenn": "university of pennsylvania",
        "uchicago": "university of chicago",
        "tokyo tech": "tokyo institute of technology",
        "uofm": "university of michigan",
        "utexas": "university of texas",
        "nyu": "new york university",
        "ncsu": "north carolina state university",
        "wustl": "washington university in st. louis",
        "duke": "duke university",
        "cornell": "cornell university",
        "berkeley": "university of california berkeley",
        "brown": "brown university",
        "purdue": "purdue university",
        "georgia tech": "georgia institute of technology",
        "uci": "university of california irvine"
    }
    # 2) common typos we want to auto-fix
    typo_fixes = {
        "univeristy": "university",
        "colledge":   "college",
        "centeres":   "centers",
        "scool":      "school",
        "instutute":  "institute",
        "laborotory": "laboratory",
    }

    # Title-Case regex: catches "Rutgers University" or "University of ... Technology"
    title_regex = re.compile(
        r'([A-Z][\w&\.\,\s]{2,50}\b(?:University|College|Institute|School|Centre|Center|Laboratory)\b)'
    )
    # fallback keyword
    keyword_re = re.compile(
        r'\b(university|college|institute|school|centre|center|laboratory)\b',
        re.IGNORECASE
    )

    file_object.seek(0)
    pdf = pymupdf.open(stream=file_object.read(), filetype="pdf")
    try:
        for page in pdf:
            text = page.get_text("text")
            # only profile pages
            if not all(marker in text for marker in ("Citations", "h-index", "i10-index")):
                continue

            lines = text.splitlines()
            header_lines = lines[:50]  # only the top-of-page block

            # === First pass: Title-Case + multi-line glue ===
            done = False
            for idx, L in enumerate(header_lines):
                m = title_regex.search(L)
                if not m:
                    continue

                inst = m.group(1).strip()
                # if the line ends with a comma (or the regex ends right before one), glue the next line
                after = L[m.end():].lstrip()
                if L.rstrip().endswith(",") or after.startswith(","):
                    if idx + 1 < len(header_lines):
                        nxt = header_lines[idx + 1].strip()
                        if nxt and nxt[0].isalpha() and nxt[0].isupper():
                            inst = f"{inst}, {nxt}"
                institutions.add(inst)
                done = True
                break
            if done:
                continue

            # === Second pass: typo-fix → alias → keyword ===
            for L in header_lines:
                nl = L.strip().lower()
                # apply typo fixes
                for bad, good in typo_fixes.items():
                    nl = nl.replace(bad, good)
                # normalize internal dashes/spaces
                nl = re.sub(r'(?<=\S)-(?=\S)', ' ', nl)
                nl = re.sub(r'[\s–—]+', ' ', nl).strip()
                # strip trailing punctuation
                nl = nl.rstrip('.,;:')

                tokens = re.findall(r'\w+', nl)

                # 2a) alias match
                matched = False
                for key, full in alias_mapping.items():
                    if key in tokens:
                        institutions.add(full)
                        matched = True
                        break
                if matched:
                    break

                # 2b) generic keyword
                if keyword_re.search(nl):
                    institutions.add(nl.title())
                    break

        return institutions

    finally:
        pdf.close()

#--------------
def find_similar_titles(set1, set2, threshold=90):
    matches = set()
    for t1 in set1:
        for t2 in set2:
            if fuzz.ratio(t1, t2) >= threshold:
                matches.add(t1)
    return matches

def find_common_institutions(set1, set2):
    return set1 & set2

@st.cache_data
def convert_df(df):
    return df.to_csv(index=True).encode('utf-8')

def create_word_file(extracted_institutions):
    doc = Document()
    doc.add_heading('Extracted Institutions', 0)
    for filename, institutions in extracted_institutions.items():
        doc.add_heading(f'Institutions from {filename}', level=1)
        for institution in institutions:
            doc.add_paragraph(institution)

    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)
    return word_stream



def extract_collaborator_names(pdf_files):
    all_names = {}

    non_name_terms = {
        "university", "college", "institute", "center", "school", "engineering", 
        "affiliation", "department", "chemical", "powder", "advanced", "jersey", 
        "active", "organization", "institution", "last", "date", "of", "research"
    }

    for i, file in enumerate(pdf_files):
        file.seek(0)
        doc = pymupdf.open(stream=file.read(), filetype="pdf")
        file_name = getattr(file, "name", f"File_{i+1}")
        names = set()

        try:
            for page in doc:
                text = page.get_text("text")
                lines = text.splitlines()

                for line in lines:
                    line = line.strip().lower()
                    if not line or any(term in line for term in non_name_terms):
                        continue

                    # Comma-separated: Last, First (e.g. "Braatz, Richard")
                    if match := re.match(r'^(?:[GT]:)?\s*([\w\-]+),\s*([\w\-]+)', line):
                        last, first = match.groups()
                        names.add(f"{first} {last}")
                        continue

                    # # G: First Last / T: First Last (e.g. "G: Richard Braatz")
                    # if match := re.match(r'^[GTCAREB]:\s*([a-z\-]+) ([a-z\-]+)', line, re.IGNORECASE):
                    #     names.add(f"{match.group(1)} {match.group(2)}")
                    #     continue
                
                   # Match lines like "G: First Last", "T First Last", "R: First Last", etc.
                    if match := re.match(r'^[GTCAREB]:?\s+([a-z\-]+)\s+([a-z\-]+)', line, re.IGNORECASE):
                       names.add(f"{match.group(1)} {match.group(2)}")
                       continue

                    

                    # Plain First Last (e.g. "Richard Braatz")
                    if match := re.match(r'^([a-z\-]+) ([a-z\-]+)$', line):
                        names.add(f"{match.group(1)} {match.group(2)}")
                        continue

                    # Handle cases like "Richard D. Braatz", "Ashley Su Wen Dan"
                    # Match: First (optional middle initial(s) or name) Last
                    if match := re.match(r'^([a-z\-]+)(?: [a-z\-]+)* ([a-z\-]+)$', line):
                        first, last = match.groups()
                        names.add(f"{first} {last}")

                    # Allow middle initial like "D." (e.g. "Richard D. Braatz" should return "Richard Braatz")
                    if match := re.match(r'^([a-z\-]+)(?: [a-z]\.)? ([a-z\-]+)$', line):
                        first, last = match.groups()
                        names.add(f"{first} {last}")

        finally:
            doc.close()

        all_names[file_name] = list(names)

    return all_names

#-------------------------------------------------------------
# def extract_researcher_name(file_object):
#     file_object.seek(0)
#     doc = pymupdf.open(stream=file_object.read(), filetype="pdf")
#     name = ""
    
#     try:
#         page = doc[0]
#         text = page.get_text("text")

#         # Search for the name with 'Google Scholar' anywhere in the text
#         for line in text.splitlines():
#             if "Google Scholar" in line:
#                 name = line.split("Google Scholar")[0].strip()
#                 break  # Exit the loop once the name is found

#     finally:
#         doc.close()

#     # Remove titles like "Dr.", "Prof.", etc.
#     name = name.lower().replace("dr. ", "").replace("prof. ", "").replace(" -", "").strip()

#     # Remove middle initials (with or without a period) - Adjusted regex
#     name = re.sub(r'(?<=\s)[a-z]\.?(\s+|$)', ' ', name)  # Removes middle initials only

#     # Remove any remaining middle names (longer names after the first and last name)
#     #name = re.sub(r'\s+[a-z]+(\s+[a-z]+)*$', '', name)  # Removes middle names (e.g., "richard d. braatz" -> "richard braatz")

#     # Ensure name is not empty and return
#     return name if name else ""

import unicodedata

def extract_researcher_name(file_object):
    file_object.seek(0)
    doc = pymupdf.open(stream=file_object.read(), filetype="pdf")
    name = ""

    try:
        page = doc[0]
        text = page.get_text("text")

        for line in text.splitlines():
            if "Google Scholar" in line:
                name = line.split("Google Scholar")[0].strip()
                break
    finally:
        doc.close()

    def clean_name(name):
        # Normalize weird/invisible characters
        name = unicodedata.normalize("NFKD", name)
        name = re.sub(r"[^\w\s\-.]", "", name)  # remove anything not letter, space, dash, dot

        # Remove academic titles
        name = name.lower().replace("dr. ", "").replace("prof. ", "").replace(" -", "").strip()

        # Remove middle initials (e.g. M.Y. or M. Y.)
        name = re.sub(r"\b[a-zA-Z]\.(?:\s*[a-zA-Z]\.)*", "", name)

        # Collapse multiple spaces into one
        name = re.sub(r"\s+", " ", name).strip()

        return name

    return clean_name(name) if name else ""


# === UI ===

st.markdown(
    "<h2 style='font-size: 40px;'>ConflictMatch: PI-Reviewer Conflicts via Common Publications, Institutions & Collaborations/Affiliations</h2>", 
    unsafe_allow_html=True
)
st.markdown('"This ConflictMatch tool is implemented based on the work by Rohit Ramachandran, Rutgers University, NJ, USA"')

st.markdown("""
### Instructions:
1. Upload PI and Reviewer Google Scholar PDFs, and PI Collaboration & other Affiliation (CoA) PDFs.
2. Ensure the full Google Scholar profile is printed (scroll down and click 'show more' in the URL.You may have to do this multiple times.)
3. Select publication years to filter. 4 years is default for COIs.
4. Click "Run Simulation" to identify publication, institution and CoA overlaps and potential conflicts.
5. View and/or Download results 
""")

current_year = datetime.datetime.now().year
year_options = ["All Years"] + list(range(current_year, 1950, -1))
selected_years = st.multiselect("Select years to include:", year_options, default=year_options[1:6])
if "All Years" in selected_years:
    selected_years = set(map(str, range(current_year, 1950, -1)))
else:
    selected_years = set(map(str, [y for y in selected_years if y != "All Years"]))
    

#-----------optional merge section for multiple PIs in the same PI Team


# Function to merge PDFs using PyMuPDF
def merge_pdfs_pymupdf(pdf_list):
    # Create a new PDF document to store the merged content
    merged_pdf = pymupdf.open()
    
    for pdf_file in pdf_list:
        # Open each uploaded PDF (pdf_file is a file-like object)
        pdf_doc = pymupdf.open(stream=pdf_file.read(), filetype="pdf")
        
        # Append each page to the merged PDF document
        merged_pdf.insert_pdf(pdf_doc)
    
    # Save the merged PDF to a BytesIO object to serve it in memory
    merged_pdf_io = io.BytesIO()
    merged_pdf.save(merged_pdf_io)
    merged_pdf_io.seek(0)  # Reset pointer to the beginning
    
    return merged_pdf_io

# Optional: File uploader for merging PDFs
merged_pdf_files = st.file_uploader(
    "STEP 0: OPTIONAL:  Upload individual Google Scholar PDFs to Merge (e.g., from a multi-investigator team). Then download, drag and drop merged files to STEP 1. Remove any uploaded files before a new merge", 
    type=["pdf"], 
    accept_multiple_files=True, 
    key="merge_uploader"
)


if merged_pdf_files:
    if len(merged_pdf_files) < 2:
        st.warning("Please upload at least **2 PDF files** to perform a merge.")
    else:
        # Ensure files are passed as a list, even if only one file is uploaded
        merged_pdf = merge_pdfs_pymupdf(merged_pdf_files)
        first_filename = merged_pdf_files[0].name

        # Provide a download link for the merged PDF
        st.download_button(
            label="⬇️ Download Merged PDF",
            data=merged_pdf,
            file_name=f"merged_{first_filename}",
            mime="application/pdf"
        )


#-----------end of optiional merge 

group1_files = st.file_uploader("STEP 1: REQUIRED: Upload Google Scholar PDFs for each PI or investigator team", type=["pdf"], accept_multiple_files=True)


#---------------------------------------------------------------------


group2_files = st.file_uploader("STEP 2: REQUIRED: Upload Google Scholar PDFs for each Reviewer", type=["pdf"], accept_multiple_files=True)

 
# Optional upload for Collaborator List files (CoA files)
collaborator_list_files = st.file_uploader(
    "STEP 3: OPTIONAL: Upload Collaboration and other Affiliations (CoA) PDFs for each PI or PI team. "
    "Note: CoA files from eJ are already merged for multi-investigators.",
    type=["pdf"],
    accept_multiple_files=True,
    key="collaborator_uploader"
)

# Check if collaborator list files are uploaded and match the length of group1_files
if collaborator_list_files and len(collaborator_list_files) != len(group1_files):
    st.error(f"The number of CoA files must match the number of Group 1 PI files. "
             f"Expected {len(group1_files)} CoA file(s), but received {len(collaborator_list_files)}.")
    st.stop()  # Stop further execution if lengths don't match


if st.button("Run Simulation"):
    # Initialize session state for storing data
    if 'simulation_data' not in st.session_state:
        st.session_state.simulation_data = {}

    # Start timing at the very beginning
    start_time = time.time()

    group1_data, group2_data = {}, {}
    group1_institutions, group2_institutions = {}, {}
    extracted_institutions = {}
    
    # Track if any publications were found
    any_publications_found = False
    
    # Check Group 1 (PIs)
    for file in group1_files:
        name = clean_filename(file.name)
        titles = extract_titles_and_years(file)
        filtered = {t for t, y in titles if y in selected_years}
        group1_data[name] = filtered
        if filtered:
            any_publications_found = True
    
    # Check Group 2 (Reviewers)
    for file in group2_files:
        name = clean_filename(file.name)
        titles = extract_titles_and_years(file)
        filtered = {t for t, y in titles if y in selected_years}
        group2_data[name] = filtered
        if filtered:
            any_publications_found = True

    # If no publications were found in any file, show error and stop
    if not any_publications_found:
        st.error("❌ No publications were found in any of the uploaded files. Please check that your PDFs are not corrupt and contain Google Scholar profile data.")
        st.stop()

    # Display extracted titles
    # st.subheader("📚 Extracted Publications")
    # for file in group1_files:
    #     name = clean_filename(file.name)
    #     titles = extract_titles_and_years(file)
    #     filtered = {t for t, y in titles if y in selected_years}
    #     st.write(f"**{name}**")
    #     for title in sorted(filtered):
    #         st.write(f"- {title}")
    #     st.write("---")

    # for file in group2_files:
    #     name = clean_filename(file.name)
    #     titles = extract_titles_and_years(file)
    #     filtered = {t for t, y in titles if y in selected_years}
    #     st.write(f"**{name}**")
    #     for title in sorted(filtered):
    #         st.write(f"- {title}")
    #     st.write("---")

    # Continue with existing code...
    group1_institutions = {name: extract_institution_names(file) for name, file in zip(group1_data.keys(), group1_files)}
    group2_institutions = {name: extract_institution_names(file) for name, file in zip(group2_data.keys(), group2_files)}
    
    # Update extracted_institutions with both group1 and group2 institutions
    extracted_institutions.update(group1_institutions)
    extracted_institutions.update(group2_institutions)

    group1_names = list(group1_data.keys())
    group2_names = list(group2_data.keys())

    # === Table 1: Common Publication Count ===
    comparison_matrix = []
    common_publications = {}

    for r1 in group1_names:
        row = []
        for r2 in group2_names:
            commons = find_similar_titles(group1_data[r1], group2_data[r2], threshold=90)
            row.append(len(commons))
            common_publications[(r1, r2)] = list(commons)
        comparison_matrix.append(row)

    df_comparison = pd.DataFrame(comparison_matrix, index=[f"PI: {n}" for n in group1_names], columns=[f"Rev: {n}" for n in group2_names])
    
    # Store in session state
    st.session_state.simulation_data['df_comparison'] = df_comparison
    st.session_state.simulation_data['common_publications'] = common_publications
    
    # Convert common_publications to DataFrame for display and storage
    # Create a matrix-style DataFrame for common publications
    pub_matrix = []
    for r1 in group1_names:
        row = []
        for r2 in group2_names:
            commons = common_publications.get((r1, r2), [])
            # Add numbers to each publication
            numbered_commons = [f"{i+1}. {pub}" for i, pub in enumerate(commons)]
            row.append("\n".join(numbered_commons) if commons else "")
        pub_matrix.append(row)
    
    common_publications_df = pd.DataFrame(
        pub_matrix,
        index=[f"PI: {n}" for n in group1_names],
        columns=[f"Rev: {n}" for n in group2_names]
    )
    st.session_state.simulation_data['common_publications_df'] = common_publications_df

    # === Table 2: COI - Publications ===
    df_pub_coi = df_comparison.applymap(lambda x: "COI" if x > 0 else "")
    
    # Store in session state
    st.session_state.simulation_data['df_pub_coi'] = df_pub_coi

    # === Table 3: COI - Institutions ===
    inst_matrix = []
    for r1 in group1_names:
        row = []
        for r2 in group2_names:
            commons = find_common_institutions(group1_institutions.get(r1, set()), group2_institutions.get(r2, set()))
            row.append("COI" if commons else "")
        inst_matrix.append(row)

    df_inst_coi = pd.DataFrame(inst_matrix, index=[f"PI: {n}" for n in group1_names], columns=[f"Rev: {n}" for n in group2_names])
    
    # Store in session state
    st.session_state.simulation_data['df_inst_coi'] = df_inst_coi

    # === Table 4: COI - Collaborations ===
    if not collaborator_list_files:
        with open("blank.pdf", "rb") as f:
            blank_data = f.read()
            collaborator_list_files = []
            for i in range(len(group1_names)):
                fake_file = io.BytesIO(blank_data)
                fake_file.name = f"blank_{i+1}.pdf"
                collaborator_list_files.append(fake_file)

    collaborator_names = extract_collaborator_names(collaborator_list_files)
    group2_researcher_names = {}
    
    for file in group2_files:
        name_key = clean_filename(file.name)
        researcher_name = extract_researcher_name(file) 
        group2_researcher_names[name_key] = researcher_name

    # Create two matrices: one for display (COI) and one for download (names)
    collab_matrix_display = []
    collab_matrix_download = []

    def is_fuzzy_match(name1, name2_list, threshold=80):
        for c in name2_list:
            if fuzz.ratio(name1, c) >= threshold:
                return True
        return False

    for i in range(len(group1_names)):
        row_display = []
        row_download = []
        for j in range(len(group2_names)):
            researcher_name = group2_researcher_names.get(group2_names[j], "")
            collab_list_file_name = getattr(collaborator_list_files[i], "name", f"File_{i+1}")
            collaborators = collaborator_names.get(collab_list_file_name, [])
            
            # Find matching collaborators
            matching_collabs = [c for c in collaborators if fuzz.ratio(researcher_name, c) >= 80]
            if matching_collabs:
                row_display.append("COI")
                row_download.append("\n".join(matching_collabs))
            else:
                row_display.append("")
                row_download.append("")
        
        collab_matrix_display.append(row_display)
        collab_matrix_download.append(row_download)

    df_collab_coi = pd.DataFrame(collab_matrix_display, index=[f"PI: {n}" for n in group1_names], columns=[f"Rev: {n}" for n in group2_names])
    df_collab_names = pd.DataFrame(collab_matrix_download, index=[f"PI: {n}" for n in group1_names], columns=[f"Rev: {n}" for n in group2_names])
    
    # Store in session state
    st.session_state.simulation_data['df_collab_coi'] = df_collab_coi
    st.session_state.simulation_data['df_collab_names'] = df_collab_names
    st.session_state.simulation_data['collaborator_names'] = collaborator_names
    st.session_state.simulation_data['group2_researcher_names'] = group2_researcher_names

    # === Combined COI Matrix ===
    combined_dict = {}

    for i, r1 in enumerate(group1_names):
        for j, r2 in enumerate(group2_names):
            pub_flag = df_pub_coi.iloc[i, j] == "COI"  
            inst_flag = df_inst_coi.iloc[i, j] == "COI"  
            collab_flag = df_collab_coi.iloc[i, j] == "COI"  
            co_flag = "COI" if pub_flag or inst_flag or collab_flag else ""
            combined_dict[(f"Grp 1: {r1}", f"Grp 2: {r2}")] = co_flag

    rows = []
    for r1 in group1_names:
        row = []
        for r2 in group2_names:
            row.append(combined_dict.get((f"Grp 1: {r1}", f"Grp 2: {r2}"), ""))
        rows.append(row)

    df_combined_coi = pd.DataFrame(
        rows,
        index=[f"PI: {n}" for n in group1_names],
        columns=[f"Rev: {n}" for n in group2_names]
    )
    
    # Store in session state
    st.session_state.simulation_data['df_combined_coi'] = df_combined_coi
    st.session_state.simulation_data['extracted_institutions'] = extracted_institutions

    # Create Word document for names
    doc = Document()
    if collaborator_names:
        doc.add_heading("📑 Extracted Collaborator Names", level=2)
        for filename in collaborator_names:
            doc.add_paragraph("")
            doc.add_paragraph(f"Source file: {filename}", style='Normal')
            for name in collaborator_names[filename]:
                doc.add_paragraph(f"• {name}", style='List Bullet')

    if group2_researcher_names:
        doc.add_heading("📑 Extracted Researcher Names (Group 2)", level=2)
        for name in group2_researcher_names.values():
            doc.add_paragraph(name)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Store in session state
    st.session_state.simulation_data['names_doc_buffer'] = buffer

    # Create Word document for institutions
    word_file = create_word_file(extracted_institutions)
    
    # Store in session state
    st.session_state.simulation_data['institutions_doc_buffer'] = word_file

    # Create pairwise common institutions data
    pairwise_common_institutions = {}
    for r1 in group1_names:
        for r2 in group2_names:
            commons = find_common_institutions(group1_institutions.get(r1, set()), group2_institutions.get(r2, set()))
            pairwise_common_institutions[(r1, r2)] = list(commons) if commons else []

    # Create a matrix-style DataFrame for common institutions
    inst_matrix = []
    for r1 in group1_names:
        row = []
        for r2 in group2_names:
            commons = pairwise_common_institutions.get((r1, r2), [])
            row.append("\n".join(commons) if commons else "")
        inst_matrix.append(row)
    
    common_institutions_df = pd.DataFrame(
        inst_matrix,
        index=[f"PI: {n}" for n in group1_names],
        columns=[f"Rev: {n}" for n in group2_names]
    )
    
    # Store in session state
    st.session_state.simulation_data['common_institutions_df'] = common_institutions_df

    # Calculate and store final simulation time
    end_time = time.time()
    simulation_time = (end_time - start_time) / 60
    st.session_state.simulation_data['simulation_time'] = simulation_time

# Display tables from session state if they exist
if 'simulation_data' in st.session_state:
    # === Table 1: Common Publication Count ===
    st.subheader("📊 Common Publication Count Table")
    st.dataframe(st.session_state.simulation_data['df_comparison'].style.set_properties(**{'text-align': 'center'}))

    # === Table 2: COI - Publications ===
    st.subheader("🔵 COI Matrix: Based on Common Co-authored Publications")
    st.dataframe(st.session_state.simulation_data['df_pub_coi'].style.applymap(lambda v: "background-color: #FFCCCC" if v == "COI" else ""))

    # === Table 3: COI - Institutions ===
    st.subheader("🟢 COI Matrix: Based on Common Institutions")
    st.dataframe(st.session_state.simulation_data['df_inst_coi'].style.applymap(lambda v: "background-color: #CCFFCC" if v == "COI" else ""))

    # === Table 4: COI - Collaborations ===
    st.subheader("🔴 COI Matrix: Based on Collaborations & Other Affiliations")
    st.dataframe(st.session_state.simulation_data['df_collab_coi'].style.applymap(lambda v: "background-color: #FFB6C1" if v == "COI" else ""))

    # === Combined COI Matrix ===
    st.subheader("🟡 Combined COI Matrix (Publications, Institutions, or Collaborations)")
    st.dataframe(st.session_state.simulation_data['df_combined_coi'].style.applymap(
        lambda v: "background-color: #FFF3B0" if v == "COI" else "")
    )

    st.subheader(f"✅ Total simulation time: {st.session_state.simulation_data['simulation_time']:.2f} minutes")

    st.subheader("📥 Download Results")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.download_button(
            "⬇️ Download Common Publication Count CSV",
            convert_df(st.session_state.simulation_data['df_comparison']),
            "common_publication_counts.csv",
            "text/csv"
        )
        
        st.download_button(
            "⬇️ Download COI (Publications) CSV",
            convert_df(st.session_state.simulation_data['df_pub_coi']),
            "coi_publications.csv",
            "text/csv"
        )
        
        st.download_button(
            "⬇️ Download COI (Institutions) CSV",
            convert_df(st.session_state.simulation_data['df_inst_coi']),
            "coi_institutions.csv",
            "text/csv"
        )
        
        st.download_button(
            "⬇️ Download COI (Collaborations) CSV",
            convert_df(st.session_state.simulation_data['df_collab_coi']),
            "coi_collaborations.csv",
            "text/csv"
        )
        
        st.download_button(
            "⬇️ Download Collaborator Names CSV",
            convert_df(st.session_state.simulation_data['df_collab_names']),
            "collaborator_names.csv",
            "text/csv"
        )
        
        st.download_button(
            "⬇️ Download Combined COI Matrix CSV",
            convert_df(st.session_state.simulation_data['df_combined_coi']),
            "coi_combined.csv",
            "text/csv"
        )
    
    with col2:
        st.download_button(
            "⬇️ Download Pairwise Common Publications",
            st.session_state.simulation_data['common_publications_df'].to_csv(index=True),
            "pairwise_common_publications.csv",
            "text/csv"
        )
        
        st.download_button(
            "⬇️ Download Pairwise Common Institutions",
            st.session_state.simulation_data['common_institutions_df'].to_csv(index=True),
            "pairwise_common_institutions.csv",
            "text/csv"
        )
        
        st.download_button(
            "⬇️ Download Extracted Names (Word Doc)",
            st.session_state.simulation_data['names_doc_buffer'],
            "extracted_names.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        st.download_button(
            "⬇️ Download Extracted Institutions (Word File)",
            st.session_state.simulation_data['institutions_doc_buffer'],
            "extracted_institutions.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )